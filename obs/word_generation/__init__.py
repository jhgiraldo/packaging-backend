import azure.functions as func
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
import logging

# =========================
#   Estilos: Montserrat
# =========================
def _apply_font_to_style(style, font_name="Montserrat", size_pt=None, bold=None):
    """Aplica Montserrat (y opcionalmente tamaño/negrita) a un estilo de Word."""
    if style is None:
        return
    style.font.name = font_name
    rpr = style.element.rPr
    if rpr is not None and rpr.rFonts is not None:
        rpr.rFonts.set(qn("w:ascii"), font_name)
        rpr.rFonts.set(qn("w:hAnsi"), font_name)
        rpr.rFonts.set(qn("w:cs"), font_name)
        rpr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold

def _ensure_montserrat_styles(doc: Document, font_name="Montserrat"):
    """Fuerza Montserrat en estilos comunes y define tamaños típicos."""
    try:
        _apply_font_to_style(doc.styles["Normal"], font_name=font_name, size_pt=11, bold=False)
        doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    except Exception:
        pass
    try:
        _apply_font_to_style(doc.styles["Title"], font_name=font_name, size_pt=24, bold=True)
    except Exception:
        pass
    for st_name, size in [("Heading 1", 16), ("Heading 2", 14), ("List Bullet", 11), ("List Paragraph", 11)]:
        try:
            _apply_font_to_style(doc.styles[st_name], font_name=font_name, size_pt=size)
        except Exception:
            pass

# =========================
#   Helpers de contenido
# =========================
def _as_text(value):
    if value is None:
        return ""
    if isinstance(value, dict):
        return value.get("item") or value.get("descripcion") or str(value)
    return str(value)

def _add_item_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p

def _add_subitem_dash(doc: Document, text: str, indent_cm: float = 0.75):
    p = doc.add_paragraph(style="List Paragraph")
    p.add_run(f"- {text}")
    p.paragraph_format.left_indent = Cm(indent_cm)
    return p

def _render_items_with_subitems(doc: Document, entries: list, *, expect_item_key: bool = True, indent_cm: float = 0.75, legacy_pair: bool = False):
    if not entries:
        return
    for e in entries:
        if legacy_pair:
            desc = _as_text(e.get("descripcion"))
            estado = _as_text(e.get("estado") or "N/A")
            _add_item_bullet(doc, f"{desc} – Estado: {estado}")
            continue
        item_text = _as_text(e.get("item") if isinstance(e, dict) else e)
        _add_item_bullet(doc, item_text)
        subitems = []
        if isinstance(e, dict):
            raw = e.get("subitems") or []
            for si in raw:
                subitems.append(_as_text(si))
        for si_text in subitems:
            _add_subitem_dash(doc, si_text, indent_cm=indent_cm)

# =========================
#   Constructor del DOCX
# =========================
def _build_doc(payload, owner="Jhonatan Giraldo Cardona") -> bytes:
    doc = Document()
    _ensure_montserrat_styles(doc, font_name="Montserrat")

    # Título dinámico
    titulo = "Reporte Comercial"
    if isinstance(payload, dict):
        titulo = payload.get("titulo", titulo)
    doc.add_heading(titulo, 0)

    clientes = payload if isinstance(payload, list) else payload.get("clientes", [])
    for c in clientes:
        nombre = (c.get("cliente") or c.get("nombre") or "Cliente")
        doc.add_heading(nombre, level=1)

        # ---- Oportunidades ----
        doc.add_heading("Oportunidades", level=2)
        oportunidades = c.get("oportunidades") or []
        legacy_mode = False
        if oportunidades and isinstance(oportunidades[0], dict) and "descripcion" in oportunidades[0]:
            legacy_mode = True
        if oportunidades:
            _render_items_with_subitems(doc, oportunidades, legacy_pair=legacy_mode, indent_cm=0.75)
        else:
            doc.add_paragraph("Sin oportunidades registradas.")

        # ---- Riesgos y bloqueos ----
        doc.add_heading("Riesgos y bloqueos", level=2)
        riesgos = c.get("riesgos") or []
        if riesgos:
            _render_items_with_subitems(doc, riesgos, indent_cm=0.75)
        else:
            doc.add_paragraph("Sin riesgos registrados.")

        # ---- Tareas pendientes ----
        doc.add_heading("Tareas Pendientes", level=2)
        tareas = c.get("tareas") or []
        if tareas:
            _render_items_with_subitems(doc, tareas, indent_cm=0.75)
        else:
            doc.add_paragraph("Sin tareas pendientes.")

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream.getvalue()

# =========================
#   Azure Function entry
# =========================
def main(req: func.HttpRequest) -> func.HttpResponse:
    logging.info("HTTP trigger -> generación de Word con bullets y subitems con guiones.")

    if req.method == "GET":
        return func.HttpResponse("OK. Envía un POST con JSON para generar el Word.", status_code=200)

    try:
        payload = req.get_json()
    except ValueError:
        return func.HttpResponse("Cuerpo JSON inválido.", status_code=400)

    owner = (
        req.params.get("owner")
        or (payload.get("owner") if isinstance(payload, dict) else None)
        or "Jhonatan Giraldo Cardona"
    )

    try:
        content = _build_doc(payload, owner=owner)
    except Exception as e:
        logging.exception("Error generando el documento.")
        return func.HttpResponse(f"Error generando el documento: {e}", status_code=500)

    # Nombre dinámico del archivo
    titulo_archivo = "ReporteComercial"
    if isinstance(payload, dict):
        titulo_archivo = payload.get("titulo", titulo_archivo).replace(" ", "")
    filename = f"{titulo_archivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

    return func.HttpResponse(
        content,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )